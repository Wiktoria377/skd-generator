"""
SKD Lawsuit Document Generator v4.
Uses HARDCODED SEQUENTIAL POSITIONAL MAPPING derived from the golden standard.
Each placeholder in the template is mapped to a specific CreditData field
by its sequential index (0-104), NOT by paragraph keyword matching.
"""
import re
from dataclasses import fields
from typing import Optional

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from app.models.schema import CaseContext, CreditData, REQUIRED_DOC_ALERTS, DocumentType


MISSING_MARKER = "[DO UZUPEŁNIENIA]"

# ============================================================
# SEQUENTIAL PLACEHOLDER MAP
# Reverse-engineered from the Brzeziński golden standard.
# Index = the Nth placeholder in the template (0-based).
# Value = CreditData field name to fill it with.
# ============================================================

PLACEHOLDER_MAP = {
    # === NAGŁÓWEK (header) ===
    0: 'data_pozwu',                        # Łódź, dnia ___ r.

    # === SĄD ===
    1: 'sad',                               # Sąd ___
    2: 'sad_wydzial',                       # wydział (dots)
    3: 'pozwany_adres',                     # adres sądu (dots)

    # === POWÓD (plaintiff = kredytobiorca ALWAYS) ===
    4: 'powod_imie_nazwisko',               # Powód: ___
    5: 'powod_adres',                       # adres powoda
    6: 'powod_pesel',                       # PESEL powoda

    # === POZWANY (defendant = bank ALWAYS) ===
    7: 'pozwany_nazwa',                     # Pozwany: ___
    8: 'pozwany_adres',                     # adres banku
    9: 'pozwany_krs',                       # KRS banku

    # === WPS / OPŁATA ===
    10: 'wps',                              # WPS: ___
    11: 'oplata',                           # Opłata: ___

    # === PETITUM (paragraph 24 — 12 placeholders in one paragraph) ===
    12: 'pozwany_nazwa',                    # od pozwanego ___ z siedzibą
    13: 'pozwany_siedziba',                 # z siedzibą w ___
    14: 'powod_imie_nazwisko',              # na rzecz powoda ___
    15: 'kwota_roszczenia',                 # kwoty ___ złotych
    16: 'kwota_roszczenia_slownie',         # (słownie: ___)
    17: 'data_wymagalnosci',                # od dnia ___ do dnia zapłaty
    18: 'numer_umowy',                      # Umowy ___ z dnia
    19: 'data_zawarcia_umowy',              # z dnia ___ zawartej
    20: 'powod_imie_nazwisko',              # pomiędzy ___ a
    21: 'pozwany_nazwa',                    # a ___
    22: 'okres_odsetek_do',                 # jedynie na dzień ___
    23: 'okres_odsetek_do',                 # po dniu ___

    # === USTALENIE (paragraph 25) ===
    24: 'numer_umowy',                      # z pozwanym (___) w dniu
    25: 'data_zawarcia_umowy',              # w dniu ___ nr
    26: 'numer_umowy',                      # nr ___

    # === WNIOSEK DOWODOWY (paragraphs 40) ===
    27: 'numer_umowy',                      # tytułem Umowy ___ z dnia
    28: 'data_zawarcia_umowy',              # z dnia ___

    # === ZAŁOŻENIA BIEGŁEGO - koszty (paragraph 44) ===
    29: 'prowizja',                         # prowizja w kwocie ___
    30: 'ubezpieczenie',                    # ubezpieczeniowe w kwocie ___

    # === ZAŁOŻENIE NR 1 (paragraph 46) ===
    31: 'kwota_pozyczki',                   # czyli wnosi ___ zł (Ck = kwota pożyczki)
    32: 'prowizja',                         # prowizji w wysokości ___ zł
    33: 'ubezpieczenie',                    # ubezpieczenia w wysokości ___ zł
    34: 'harmonogram_raty_opis',            # raty ... czyli ___

    # === ZAŁOŻENIE NR 2 (paragraph 47) ===
    35: 'calkowita_kwota_kredytu',          # wynosi ___ zł (Ck = CKK)
    36: 'harmonogram_raty_opis',            # raty ... czyli ___

    # === ZAŁOŻENIE NR 3 (paragraph 48) ===
    37: 'calkowita_kwota_kredytu',          # wynosi ___ zł
    38: 'prowizja',                         # prowizji w wysokości ___ zł
    39: 'ubezpieczenie',                    # ubezpieczenia w wysokości ___ zł

    # === ZAŁOŻENIE NR 4 (paragraph 49) ===
    40: 'calkowita_kwota_kredytu',          # wynosi ___ zł

    # === ZAŁOŻENIE NR 5 (paragraph 50) ===
    41: 'calkowita_kwota_kredytu',          # wynosi ___ zł
    42: 'prowizja',                         # prowizji w wysokości ___ zł
    43: 'ubezpieczenie',                    # ubezpieczenia w wysokości ___ zł

    # === WYMAGALNOŚĆ (paragraph 55) ===
    44: 'data_wymagalnosci',                # na dzień ___ kiedy to

    # === UZASADNIENIE - FAKTY (paragraph 67) ===
    45: 'data_zawarcia_umowy',              # W dniu ___ powód zawarł
    46: 'kwota_pozyczki',                   # pożyczki w wysokości ___ zł
    47: 'kwota_pozyczki_slownie',            # (słownie: ___)
    48: 'liczba_rat',                       # w ___ ratach
    49: 'dzien_platnosci_raty',             # do ___ dnia każdego miesiąca
    50: 'typ_oprocentowania',               # według ___ stopy procentowej
    51: 'oprocentowanie',                   # wynosiła ___
    52: 'kwota_pozyczki',                   # pożyczka w kwocie ___ zł
    53: 'opis_kosztow_kredytowanych',       # tj. ___

    # === NALICZONE ODSETKI (paragraph 69) ===
    54: 'paragraf_calkowita_kwota',         # określona w ___ Umowy
    55: 'calkowita_kwota_kredytu',          # wskazana na kwotę ___ zł
    56: 'kwota_pozyczki',                   # naliczone od kwoty ___ zł
    57: 'suma_odsetek_bank',                # Suma odsetek wyniosła ___ zł
    58: 'calkowity_koszt_kredytu',          # koszt kredytu: ___ zł
    59: 'kwota_raty',                       # Rata wyniosła ___ zł
    60: 'calkowity_koszt_kredytu',          # Całkowity koszt ustalony na: ___ zł

    # === HIPOTETYCZNA RATA (paragraph 70) ===
    61: 'hipotetyczna_rata',                # rata wynosiłaby ___ zł
    62: 'hipotetyczny_calkowity_koszt',     # koszt wyniósłby: ___ zł
    63: 'hipotetyczna_calkowita_do_zaplaty', # całkowita do zapłaty ___ zł

    # === PARAGRAFY NARUSZEŃ ===
    64: 'paragraf_wczesniejsza_splata',     # W ___ Umowy wskazano (wcześniejsza spłata)
    65: 'paragraf_odstapienie',             # W ___ Umowy pozwany poinformował (odstąpienie)

    # === DATY PROCESOWE ===
    66: 'data_oswiadczenia_skd',            # W dniu ___ powód złożył oświadczenie
    67: 'data_wniosku_zaswiadczenie',       # W dniu ___ powód wystąpił
    68: 'data_wezwania_do_zaplaty',         # W dniu ___ powód wezwał

    # === KWOTA ROSZCZENIA (paragraph 81) ===
    69: 'data_zawarcia_umowy',              # uiszczona w dniu ___ (data zapłaty prowizji = data umowy)
    70: 'prowizja',                         # prowizja w kwocie ___ zł
    71: 'kwota_odsetek_zaplaconych',        # kwota odsetek ___ zł
    72: 'okres_odsetek_od',                 # za okres od ___
    73: 'okres_odsetek_do',                 # do ___
    74: 'kwota_roszczenia',                 # łącznie daje kwotę ___ zł

    # === PROWIZJA UISZCZONA (paragraph 156) ===
    75: 'paragraf_prowizja_uiszczona',      # Zgodnie z ___ Umowy prowizja

    # === RRSO KOLUMNY (paragraphs 162-164) ===
    76: 'rrso_kolumna3',                    # Kolumna 3 wynosi ___
    77: 'rrso_kolumna3_roznica',            # wyższy od bank o ___
    78: 'rrso_kolumna4',                    # Kolumna 4 wynosi ___
    79: 'rrso_kolumna4_roznica',            # niższy od bank o ___
    80: 'rrso_kolumna5',                    # Kolumna 5 wynosi ___
    81: 'rrso_kolumna5_roznica',            # wyższe od banku o ___

    # === RRSO POSTANOWIENIA (paragraph 176) ===
    82: 'paragraf_rrso',                    # vide: ___

    # === CAŁKOWITA KWOTA DO ZAPŁATY (paragraph 182) ===
    83: 'odsetki_od_kredytowanych_kosztow', # wynoszącą kwotę ___ zł
    84: 'calkowita_kwota_do_zaplaty',       # zamiast ___ zł
    85: 'calkowita_kwota_do_zaplaty_prawidlowa',  # powinna wynieść ___ zł

    # === ZMIANA OPŁAT (paragraphs 188-190) ===
    86: 'paragraf_zmiana_oplat',            # W § ___ Umowy określono
    87: 'kryteria_zmiany_oplat_1',          # kryterium 1
    88: 'kryteria_zmiany_oplat_2',          # kryterium 2

    # === WCZEŚNIEJSZA SPŁATA - § ust. (paragraph 196) ===
    89: 'paragraf_wczesniejsza_splata',     # w § ___
    90: 'paragraf_ustep_wczesniejsza_splata', # ust. ___
    91: 'paragraf_wczesniejsza_splata',     # § ___
    92: 'paragraf_ustep_wczesniejsza_splata', # ust. ___

    # === ODSTĄPIENIE - § ust. (paragraph 199) ===
    93: 'paragraf_odstapienie',             # W § ___
    94: 'paragraf_ustep_odstapienie',       # ust. ___

    # === TERMIN OŚWIADCZENIA (paragraph 202) ===
    95: 'data_oswiadczenia_skd',            # w dniu ___, powód dochował

    # === PODSUMOWANIE KOŃCOWE (paragraph 231) ===
    96: 'data_oswiadczenia_skd',            # składając w dniu ___ r.

    # === KWOTY KOŃCOWE (paragraph 233) ===
    97: 'prowizja',                         # prowizji w wysokości ___
    98: 'okres_odsetek_od',                 # od dnia ___
    99: 'okres_odsetek_do',                 # do dnia ___
    100: 'kwota_odsetek_zaplaconych',       # w wysokości ___ PLN
    101: 'kwota_roszczenia',                # daje kwotę ___ zł

    # === USTALENIE KOŃCOWE (paragraph 234) ===
    102: 'numer_umowy',                     # z pozwanym (___)
    103: 'data_zawarcia_umowy',             # w dniu ___
    104: 'numer_umowy',                     # nr ___
}


def generate_lawsuit_autonomous(
    template_path: str,
    output_path: str,
    case_context: CaseContext,
) -> tuple[str, list[dict]]:
    """Generate lawsuit using hardcoded sequential placeholder mapping."""
    doc = Document(template_path)
    report = []
    cd = case_context.credit_data
    cd_dict = {f.name: getattr(cd, f.name) for f in fields(cd)}

    # Collect all placeholders sequentially across the entire document
    all_placeholders = []
    paragraph_refs = []

    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text.strip():
            continue
        for pattern in [r'_{3,}', r'…{2,}', r'\.{4,}']:
            for m in re.finditer(pattern, text):
                all_placeholders.append(m)
                paragraph_refs.append(paragraph)

    # Now fill each placeholder by its sequential index
    # Group by paragraph to process each paragraph once
    para_placeholder_groups = {}
    for idx, (match, para) in enumerate(zip(all_placeholders, paragraph_refs)):
        para_id = id(para)
        if para_id not in para_placeholder_groups:
            para_placeholder_groups[para_id] = {'para': para, 'items': []}

        field_name = PLACEHOLDER_MAP.get(idx, 'unknown')
        value = cd_dict.get(field_name) if field_name != 'unknown' else None

        para_placeholder_groups[para_id]['items'].append({
            'index': idx,
            'match': match,
            'field': field_name,
            'value': value,
        })

        report.append({
            'placeholder_index': idx,
            'field': field_name,
            'value': value,
            'status': 'filled' if value else 'missing',
        })

    # Process each paragraph
    for para_id, group in para_placeholder_groups.items():
        _fill_paragraph(group['para'], group['items'], case_context)

    doc.save(output_path)
    return output_path, report


def _fill_paragraph(paragraph, items: list[dict], ctx: CaseContext):
    """Replace all placeholders in a paragraph with their mapped values."""
    runs = paragraph.runs
    if not runs:
        return

    full_text = ''.join(r.text for r in runs)

    # Preserve formatting from first run
    ref_bold = runs[0].bold
    ref_size = runs[0].font.size
    ref_name = runs[0].font.name

    # Find all placeholder matches in the paragraph text
    placeholder_regex = re.compile(r'_{3,}|…{2,}|\.{4,}')
    matches = list(placeholder_regex.finditer(full_text))

    if len(matches) != len(items):
        # Mismatch — fallback: try to pair them by order
        pass

    # Build replacement segments
    segments = []
    last_end = 0

    for i, match in enumerate(matches):
        # Text before this placeholder
        if match.start() > last_end:
            segments.append(('original', full_text[last_end:match.start()]))

        # Get the value for this placeholder
        if i < len(items):
            value = items[i].get('value')
            field = items[i].get('field', 'unknown')
        else:
            value = None
            field = 'unknown'

        if value:
            segments.append(('filled', str(value)))
        else:
            marker = _get_missing_marker(field, ctx)
            segments.append(('missing', marker))

        last_end = match.end()

    # Remaining text after last placeholder
    if last_end < len(full_text):
        segments.append(('original', full_text[last_end:]))

    # Clear all existing runs
    for run in runs:
        run.text = ''

    # Write segments with highlighting
    first = True
    for seg_type, seg_text in segments:
        if first and runs:
            run = runs[0]
            run.text = seg_text
            first = False
        else:
            run = paragraph.add_run(seg_text)

        # Copy formatting
        if ref_bold is not None:
            run.bold = ref_bold
        if ref_size:
            run.font.size = ref_size
        if ref_name:
            run.font.name = ref_name

        # Apply highlighting
        if seg_type == 'filled':
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif seg_type == 'missing':
            run.font.highlight_color = WD_COLOR_INDEX.RED
            run.bold = True


def _get_missing_marker(field_name: str, ctx: CaseContext) -> str:
    """Generate appropriate missing marker based on what document is missing."""
    field_to_doc = {
        'data_zawarcia_umowy': DocumentType.UMOWA,
        'numer_umowy': DocumentType.UMOWA,
        'kwota_pozyczki': DocumentType.UMOWA,
        'calkowita_kwota_kredytu': DocumentType.UMOWA,
        'oprocentowanie': DocumentType.UMOWA,
        'prowizja': DocumentType.UMOWA,
        'kwota_raty': DocumentType.UMOWA,
        'suma_odsetek_bank': DocumentType.UMOWA,
        'calkowity_koszt_kredytu': DocumentType.UMOWA,
        'calkowita_kwota_do_zaplaty': DocumentType.UMOWA,
        'suma_odsetek_zaplaconych': DocumentType.ZASWIADCZENIE,
        'suma_kapitalu_splaconego': DocumentType.ZASWIADCZENIE,
        'kwota_odsetek_zaplaconych': DocumentType.ZASWIADCZENIE,
        'data_wezwania_do_zaplaty': DocumentType.WEZWANIE,
        'data_oswiadczenia_skd': DocumentType.OSWIADCZENIE_SKD,
    }

    doc_type = field_to_doc.get(field_name)
    if doc_type and doc_type in ctx.missing_documents:
        return REQUIRED_DOC_ALERTS.get(doc_type, MISSING_MARKER)

    return MISSING_MARKER
