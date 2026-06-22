"""
Hardcore multi-format extraction engine for Polish banking documents.
Handles: PDF (digital + scanned), DOCX, Excel (.xls/.xlsx/.csv), images.
Polish diacritics-optimized OCR pipeline. Claude Vision API fallback.
"""
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

try:
    import pytesseract
    from PIL import Image, ImageFilter, ImageEnhance
    from pdf2image import convert_from_path
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

try:
    from PyPDF2 import PdfReader
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False


SUPPORTED_EXTENSIONS = {
    '.pdf', '.docx', '.doc',
    '.xlsx', '.xls', '.csv', '.tsv',
    '.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif',
}

TESSERACT_CONFIGS = [
    '--oem 3 --psm 6 -l pol',   # uniform block of text
    '--oem 3 --psm 4 -l pol',   # single column, variable sizes
    '--oem 3 --psm 3 -l pol',   # fully automatic
    '--oem 3 --psm 11 -l pol',  # sparse text - for scattered banking docs
]


class ExtractionResult:
    """Result with metadata about how text was extracted."""
    def __init__(self, text: str, method: str, page_count: int = 1,
                 tables: Optional[list] = None, confidence: float = 1.0):
        self.text = text
        self.method = method
        self.page_count = page_count
        self.tables = tables or []
        self.confidence = confidence

    @property
    def has_tables(self) -> bool:
        return len(self.tables) > 0


def extract_text(filepath: str) -> str:
    """Top-level extraction - returns plain text. Backward compatible."""
    result = extract_deep(filepath)
    combined = result.text
    if result.tables:
        combined += "\n\n=== DANE TABELARYCZNE ===\n"
        for i, tbl in enumerate(result.tables):
            combined += f"\n--- Tabela {i+1} ---\n{tbl}\n"
    return combined


def extract_deep(filepath: str) -> ExtractionResult:
    """Deep extraction with full metadata."""
    ext = Path(filepath).suffix.lower()

    if ext == '.docx':
        return _extract_from_docx(filepath)
    elif ext == '.doc':
        return _convert_doc_to_docx_and_extract(filepath)
    elif ext == '.pdf':
        return _extract_from_pdf_deep(filepath)
    elif ext in ('.xlsx', '.xls'):
        return _extract_from_excel(filepath)
    elif ext == '.csv':
        return _extract_from_csv(filepath)
    elif ext == '.tsv':
        return _extract_from_csv(filepath, sep='\t')
    elif ext in ('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif'):
        return _extract_from_image_deep(filepath)
    else:
        raise ValueError(f"Nieobsługiwany format pliku: {ext}")


# ============================================================
# DOCX
# ============================================================

def _extract_from_docx(filepath: str) -> ExtractionResult:
    if not HAS_DOCX:
        raise RuntimeError("python-docx nie jest zainstalowany")
    doc = DocxDocument(filepath)
    paragraphs = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            paragraphs.append(txt)

    tables_text = []
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(' | '.join(cells))
        table_str = '\n'.join(rows_data)
        tables_text.append(table_str)
        paragraphs.append(table_str)

    return ExtractionResult(
        text='\n'.join(paragraphs),
        method='python-docx',
        tables=tables_text,
    )


def _convert_doc_to_docx_and_extract(filepath: str) -> ExtractionResult:
    try:
        result = subprocess.run(
            ['textutil', '-convert', 'txt', '-stdout', filepath],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return ExtractionResult(text=result.stdout, method='textutil')
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    raise RuntimeError(f"Nie można skonwertować pliku .doc: {filepath}")


# ============================================================
# PDF - multi-layer extraction with fallbacks
# ============================================================

def _extract_from_pdf_deep(filepath: str) -> ExtractionResult:
    """Try every PDF extraction method, use the best result."""
    results = []

    # Method 1: PyPDF2 (fast, digital PDFs)
    text_pypdf2 = _try_pypdf2(filepath)
    if text_pypdf2 and _text_quality_score(text_pypdf2) > 0.3:
        results.append(ExtractionResult(
            text=text_pypdf2,
            method='PyPDF2',
            page_count=_count_pdf_pages(filepath),
            confidence=_text_quality_score(text_pypdf2),
        ))

    # Method 2: pdftotext with layout preservation
    text_pdftotext = _try_pdftotext(filepath)
    if text_pdftotext and _text_quality_score(text_pdftotext) > 0.3:
        results.append(ExtractionResult(
            text=text_pdftotext,
            method='pdftotext',
            page_count=_count_pdf_pages(filepath),
            confidence=_text_quality_score(text_pdftotext),
        ))

    # Method 3: pdftotext raw (no layout - sometimes better for tables)
    text_raw = _try_pdftotext_raw(filepath)
    if text_raw and _text_quality_score(text_raw) > 0.3:
        results.append(ExtractionResult(
            text=text_raw,
            method='pdftotext-raw',
            page_count=_count_pdf_pages(filepath),
            confidence=_text_quality_score(text_raw),
        ))

    # Pick the best digital extraction
    if results:
        best = max(results, key=lambda r: len(r.text) * r.confidence)
        if len(best.text.strip()) > 200 and best.confidence > 0.5:
            return best

    # Method 4: OCR fallback for scanned PDFs
    try:
        return _ocr_pdf_advanced(filepath)
    except Exception:
        # OCR failed — return the best digital text we got, even if short
        if results:
            return max(results, key=lambda r: len(r.text) * r.confidence)
        return ExtractionResult(
            text="[Nie udało się wyodrębnić tekstu z PDF]",
            method="extraction-failed",
            confidence=0.0,
            page_count=_count_pdf_pages(filepath),
        )


def _try_pypdf2(filepath: str) -> Optional[str]:
    if not HAS_PYPDF2:
        return None
    try:
        reader = PdfReader(filepath)
        texts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                texts.append(t)
        return '\n\n'.join(texts) if texts else None
    except Exception:
        return None


def _try_pdftotext(filepath: str) -> Optional[str]:
    try:
        result = subprocess.run(
            ['pdftotext', '-layout', filepath, '-'],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


def _try_pdftotext_raw(filepath: str) -> Optional[str]:
    try:
        result = subprocess.run(
            ['pdftotext', '-raw', filepath, '-'],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


def _count_pdf_pages(filepath: str) -> int:
    if HAS_PYPDF2:
        try:
            return len(PdfReader(filepath).pages)
        except Exception:
            pass
    return 1


def _text_quality_score(text: str) -> float:
    """Score 0-1 for how much the text looks like real Polish text vs garbage."""
    if not text or len(text.strip()) < 20:
        return 0.0

    total_chars = len(text)
    alpha_chars = sum(1 for c in text if c.isalpha())
    alpha_ratio = alpha_chars / total_chars if total_chars > 0 else 0

    polish_words = ['umowa', 'kredyt', 'pożyczk', 'kwota', 'rata', 'bank',
                    'odsetki', 'prowizj', 'spłat', 'oprocentow', 'złot',
                    'okres', 'termin', 'harmonogram', 'zaświadcz', 'wezwani',
                    'dnia', 'roku', 'zł', 'PLN', 'nr', 'ust.', 'art.', 'pkt']
    text_lower = text.lower()
    word_hits = sum(1 for w in polish_words if w.lower() in text_lower)
    word_score = min(word_hits / 5.0, 1.0)

    polish_diacritics = 'ąćęłńóśźżĄĆĘŁŃÓŚŹŻ'
    diac_count = sum(1 for c in text if c in polish_diacritics)
    diac_ratio = min(diac_count / max(alpha_chars, 1) * 20, 1.0)

    return (alpha_ratio * 0.3 + word_score * 0.5 + diac_ratio * 0.2)


def _ocr_pdf_advanced(filepath: str) -> ExtractionResult:
    """Advanced OCR with multiple PSM modes and image preprocessing."""
    if not HAS_OCR:
        return ExtractionResult(
            text="[SKAN PDF - brak OCR: zainstaluj tesseract i poppler]",
            method="ocr-unavailable",
            confidence=0.0,
            page_count=_count_pdf_pages(filepath),
        )

    all_texts = []
    MAX_OCR_PAGES = 10  # Never OCR more than 10 pages — key data is always in first pages
    try:
        try:
            images = convert_from_path(filepath, dpi=300, last_page=MAX_OCR_PAGES)
        except Exception:
            try:
                images = convert_from_path(filepath, dpi=200, last_page=MAX_OCR_PAGES)
            except Exception:
                images = convert_from_path(filepath, dpi=150, last_page=5)

        for i, img in enumerate(images):
            best_text = ""
            best_score = 0.0

            preprocessed_variants = _generate_image_variants(img)

            for variant in preprocessed_variants:
                for config in TESSERACT_CONFIGS[:2]:
                    try:
                        text = pytesseract.image_to_string(variant, config=config)
                        score = _text_quality_score(text)
                        if score > best_score:
                            best_score = score
                            best_text = text
                        if score > 0.7:
                            break
                    except Exception:
                        continue
                if best_score > 0.7:
                    break

            all_texts.append(f"--- Strona {i + 1} ---\n{best_text}")

        return ExtractionResult(
            text='\n\n'.join(all_texts),
            method='tesseract-advanced',
            page_count=len(images),
            confidence=_text_quality_score('\n'.join(all_texts)),
        )
    except Exception as e:
        partial = '\n\n'.join(all_texts) if all_texts else ""
        return ExtractionResult(
            text=f"{partial}\n[OCR ERROR: {e}]".strip(),
            method='tesseract-failed',
            page_count=_count_pdf_pages(filepath),
            confidence=0.1 if partial else 0.0,
        )


def _generate_image_variants(img: 'Image.Image') -> list:
    """Generate multiple preprocessed variants for OCR."""
    variants = []

    # Variant 1: High-contrast grayscale
    gray = img.convert('L')
    enhanced = ImageEnhance.Contrast(gray).enhance(2.0)
    sharp = enhanced.filter(ImageFilter.SHARPEN)
    variants.append(sharp)

    # Variant 2: Adaptive threshold (aggressive)
    thresh = gray.point(lambda x: 0 if x < 140 else 255)
    variants.append(thresh)

    # Variant 3: Lighter threshold for faded documents
    thresh_light = gray.point(lambda x: 0 if x < 100 else 255)
    variants.append(thresh_light)

    # Variant 4: Original (sometimes best for clean scans)
    variants.append(img)

    return variants


# ============================================================
# EXCEL - structured table extraction
# ============================================================

PAYMENT_HEADER_SYNONYMS = {
    'data': ['data', 'data operacji', 'data wpłaty', 'data płatności',
             'data spłaty', 'termin', 'termin płatności', 'data raty',
             'data księgowania', 'data waluty'],
    'rata_nr': ['nr', 'nr raty', 'numer', 'numer raty', 'l.p.', 'lp',
                'rata nr'],
    'kwota_raty': ['kwota', 'kwota raty', 'rata', 'kwota wpłaty', 'wpłata',
                   'kwota spłaty', 'wysokość raty', 'kwota do zapłaty'],
    'kapital': ['kapitał', 'część kapitałowa', 'spłata kapitału',
                'kapitał raty', 'rata kapitałowa', 'capital',
                'należność główna'],
    'odsetki': ['odsetki', 'część odsetkowa', 'odsetki raty',
                'rata odsetkowa', 'odsetki umowne', 'odsetki bieżące',
                'interest'],
    'prowizja': ['prowizja', 'opłata', 'opłata przygotowawcza',
                 'prowizja za udzielenie', 'commission'],
    'ubezpieczenie': ['ubezpieczenie', 'składka', 'składka ubezpieczeniowa',
                      'insurance', 'ubezp.'],
    'saldo': ['saldo', 'saldo po spłacie', 'saldo zadłużenia',
              'pozostało do spłaty', 'saldo końcowe', 'kapitał pozostały',
              'balance'],
    'typ_operacji': ['typ', 'typ operacji', 'rodzaj', 'opis', 'tytuł',
                     'opis operacji'],
}


def _extract_from_excel(filepath: str) -> ExtractionResult:
    """Deep Excel extraction with header analysis and payment timeline reconstruction."""
    if not HAS_OPENPYXL:
        raise RuntimeError("openpyxl nie jest zainstalowany: pip install openpyxl")

    wb = openpyxl.load_workbook(filepath, data_only=True)
    all_text_parts = []
    all_tables = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        all_text_parts.append(f"=== Arkusz: {sheet_name} ===")

        # Find header row
        header_row_idx, column_map = _find_header_row(rows)

        if header_row_idx is not None and column_map:
            # Structured extraction with identified headers
            table_text = _extract_structured_table(rows, header_row_idx, column_map, sheet_name)
            all_text_parts.append(table_text)
            all_tables.append(table_text)

            # Calculate summaries
            summary = _calculate_excel_summaries(rows, header_row_idx, column_map)
            if summary:
                all_text_parts.append(f"\n--- Podsumowanie obliczeniowe ({sheet_name}) ---")
                all_text_parts.append(summary)
        else:
            # Fallback: dump all non-empty rows
            for row in rows:
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    all_text_parts.append(' | '.join(cells))

    wb.close()

    return ExtractionResult(
        text='\n'.join(all_text_parts),
        method='openpyxl-structured',
        tables=all_tables,
    )


def _find_header_row(rows: list) -> tuple:
    """Find the header row by matching column names to known synonyms."""
    for row_idx, row in enumerate(rows[:15]):
        if row is None:
            continue
        cells = [str(c).strip().lower() if c is not None else '' for c in row]
        matches = {}
        for col_idx, cell_val in enumerate(cells):
            if not cell_val:
                continue
            for semantic_key, synonyms in PAYMENT_HEADER_SYNONYMS.items():
                for syn in synonyms:
                    if syn in cell_val or cell_val in syn:
                        matches[semantic_key] = col_idx
                        break

        if len(matches) >= 2:
            return row_idx, matches

    return None, {}


def _extract_structured_table(rows: list, header_idx: int, column_map: dict,
                               sheet_name: str) -> str:
    """Extract structured data using identified column mapping."""
    header_row = rows[header_idx]
    header_labels = [str(c).strip() if c else f'Kol_{i}' for i, c in enumerate(header_row)]

    lines = []
    lines.append(f"Rozpoznane kolumny: {dict((k, header_labels[v]) for k, v in column_map.items() if v < len(header_labels))}")
    lines.append(' | '.join(header_labels))
    lines.append('-' * 80)

    data_rows = rows[header_idx + 1:]
    for row in data_rows:
        if row is None:
            continue
        cells = [str(c).strip() if c is not None else '' for c in row]
        if any(c for c in cells):
            lines.append(' | '.join(cells))

    return '\n'.join(lines)


def _calculate_excel_summaries(rows: list, header_idx: int, column_map: dict) -> str:
    """Calculate sum totals from Excel payment data."""
    summaries = []
    data_rows = rows[header_idx + 1:]

    for field_name, display_name in [
        ('kwota_raty', 'Suma wpłaconych rat'),
        ('kapital', 'Suma spłaconego kapitału'),
        ('odsetki', 'Suma zapłaconych odsetek'),
        ('prowizja', 'Suma prowizji'),
        ('ubezpieczenie', 'Suma ubezpieczenia'),
    ]:
        col_idx = column_map.get(field_name)
        if col_idx is None:
            continue

        total = 0.0
        count = 0
        for row in data_rows:
            if row is None or col_idx >= len(row):
                continue
            val = row[col_idx]
            num = _parse_number(val)
            if num is not None and num != 0:
                total += num
                count += 1

        if count > 0:
            summaries.append(f"{display_name}: {total:.2f} zł (z {count} wpisów)")

    first_date = _find_first_last_date(data_rows, column_map.get('data'))
    if first_date:
        summaries.append(f"Zakres dat: {first_date}")

    return '\n'.join(summaries)


def _parse_number(val) -> Optional[float]:
    """Parse a number from various Polish formats."""
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return float(val)
    s = str(val).strip()
    s = s.replace(' ', '').replace('\xa0', '')
    s = s.replace('zł', '').replace('PLN', '').replace('złotych', '')
    s = s.strip()
    if not s:
        return None
    # Handle "1 234,56" format
    if ',' in s and '.' in s:
        s = s.replace('.', '').replace(',', '.')
    elif ',' in s:
        s = s.replace(',', '.')
    try:
        return float(s)
    except ValueError:
        return None


def _find_first_last_date(data_rows: list, date_col: Optional[int]) -> Optional[str]:
    if date_col is None:
        return None
    dates = []
    for row in data_rows:
        if row is None or date_col >= len(row):
            continue
        val = row[date_col]
        if val is not None and str(val).strip():
            dates.append(str(val).strip())
    if dates:
        return f"{dates[0]} — {dates[-1]}"
    return None


def _extract_from_csv(filepath: str, sep: str = ',') -> ExtractionResult:
    """Extract from CSV with auto-delimiter detection."""
    if not HAS_PANDAS:
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            return ExtractionResult(text=f.read(), method='raw-csv')

    for encoding in ['utf-8', 'cp1250', 'iso-8859-2', 'latin-1']:
        for delimiter in [sep, ';', '\t', ',']:
            try:
                df = pd.read_csv(filepath, encoding=encoding, sep=delimiter,
                                 engine='python', on_bad_lines='skip')
                if len(df.columns) >= 2 and len(df) >= 1:
                    text = f"Kolumny: {', '.join(str(c) for c in df.columns)}\n"
                    text += df.to_string(index=False)
                    return ExtractionResult(text=text, method='pandas-csv',
                                         tables=[text])
            except Exception:
                continue

    with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
        return ExtractionResult(text=f.read(), method='raw-fallback')


# ============================================================
# IMAGE
# ============================================================

def _extract_from_image_deep(filepath: str) -> ExtractionResult:
    if not HAS_OCR:
        raise RuntimeError("pytesseract nie jest zainstalowany")
    img = Image.open(filepath)

    best_text = ""
    best_score = 0.0

    for variant in _generate_image_variants(img):
        for config in TESSERACT_CONFIGS[:2]:
            try:
                text = pytesseract.image_to_string(variant, config=config)
                score = _text_quality_score(text)
                if score > best_score:
                    best_score = score
                    best_text = text
            except Exception:
                continue

    return ExtractionResult(
        text=best_text,
        method='tesseract-image',
        confidence=best_score,
    )
